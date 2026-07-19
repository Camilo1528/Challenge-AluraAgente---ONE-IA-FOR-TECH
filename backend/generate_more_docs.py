import os
import json
import pandas as pd
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

data_dir = "data"
os.makedirs(data_dir, exist_ok=True)

def create_pdf(filename, text_lines):
    path = os.path.join(data_dir, filename)
    c = canvas.Canvas(path, pagesize=letter)
    y = 750
    for line in text_lines:
        c.drawString(72, y, line)
        y -= 20
    c.save()

# ==========================================
# RRHH (rh, empleado, asistencia)
# Current: politicas_rh.md, manual_empleados.docx, registro_asistencia.csv
# Need: JSON, XLSX, TXT, PDF
# ==========================================

# RRHH JSON
empleados_json = [
    {"id": "E001", "nombre": "Juan Perez", "cargo": "Vendedor", "departamento": "RRHH"},
    {"id": "E002", "nombre": "Ana Gomez", "cargo": "Soporte", "departamento": "RRHH"}
]
with open(os.path.join(data_dir, "empleados_directorio.json"), "w", encoding="utf-8") as f:
    json.dump(empleados_json, f, indent=4)

# RRHH XLSX
nomina_data = {"Empleado": ["Juan Perez", "Ana Gomez"], "Salario_Base": [1200, 1300], "Bonos": [100, 150]}
pd.DataFrame(nomina_data).to_excel(os.path.join(data_dir, "empleados_nomina.xlsx"), index=False)

# RRHH TXT
with open(os.path.join(data_dir, "empleados_valores.txt"), "w", encoding="utf-8") as f:
    f.write("Valores de RRHH:\n1. Transparencia\n2. Trabajo en equipo\n3. Innovacion")

# RRHH PDF
create_pdf("empleados_beneficios.pdf", [
    "Beneficios para Empleados",
    "- Seguro de gastos medicos mayores",
    "- 30 dias de aguinaldo",
    "- Vales de despensa"
])

# ==========================================
# FINANZAS (finanzas, venta, reporte)
# Current: finanzas.json, reporte_ventas_Q2.xlsx
# Need: CSV, DOCX, TXT, PDF, MD
# ==========================================

# Finanzas CSV
gastos_data = {"Categoria": ["Marketing", "Operaciones"], "Monto": [5000, 15000], "Mes": ["Julio", "Julio"]}
pd.DataFrame(gastos_data).to_csv(os.path.join(data_dir, "reporte_gastos.csv"), index=False)

# Finanzas DOCX
doc = Document()
doc.add_heading('Reporte de Auditoria Financiera', 0)
doc.add_paragraph('La auditoria del mes de junio revela un balance positivo en todas las sucursales.')
doc.save(os.path.join(data_dir, "finanzas_auditoria.docx"))

# Finanzas TXT
with open(os.path.join(data_dir, "venta_objetivos.txt"), "w", encoding="utf-8") as f:
    f.write("Objetivos de Venta Q3:\n- Aumentar ventas online en 20%\n- Reducir costos de envio")

# Finanzas PDF
create_pdf("finanzas_balance.pdf", [
    "Balance General - Q2",
    "Activos: $150,000",
    "Pasivos: $50,000",
    "Capital: $100,000"
])

# Finanzas MD
with open(os.path.join(data_dir, "finanzas_presupuesto.md"), "w", encoding="utf-8") as f:
    f.write("# Presupuesto Anual\n\nEl presupuesto para 2026 se estima en $500,000 distribuidos equitativamente.")

# ==========================================
# INVENTARIO (inventario, proveedor, bodega)
# Current: inventario.csv, inventario.xlsx, proveedores.json
# Need: DOCX, TXT, PDF, MD
# ==========================================

# Inventario DOCX
doc = Document()
doc.add_heading('Protocolo de Bodega', 0)
doc.add_paragraph('Todo el inventario debe ser registrado en el sistema antes de ser almacenado en los estantes A y B.')
doc.save(os.path.join(data_dir, "bodega_protocolo.docx"))

# Inventario TXT
with open(os.path.join(data_dir, "inventario_alertas.txt"), "w", encoding="utf-8") as f:
    f.write("Alertas de Inventario:\n- Laptops Pro X con bajo stock (solo 5 unidades)\n- Monitores 4K agotados")

# Inventario PDF
create_pdf("proveedor_contratos.pdf", [
    "Contratos de Proveedores",
    "TechSupplies: Contrato vence en Dic 2026",
    "Global Electronics: Renovacion automatica"
])

# Inventario MD
with open(os.path.join(data_dir, "inventario_categorias.md"), "w", encoding="utf-8") as f:
    f.write("# Categorias de Inventario\n- Laptops\n- Perifericos\n- Accesorios\n- Monitores")

print("Documentos adicionales generados exitosamente en todos los formatos.")
