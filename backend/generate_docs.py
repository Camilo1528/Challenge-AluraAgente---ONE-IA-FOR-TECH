import os
import json
import pandas as pd
from docx import Document

data_dir = "data"

# 1. Crear un JSON: proveedores.json
proveedores = [
    {"id": "P001", "nombre": "TechSupplies Inc", "categoria": "Componentes", "contacto": "ventas@techsupplies.com"},
    {"id": "P002", "nombre": "Global Electronics", "categoria": "Computadoras", "contacto": "global@gelectronics.com"},
    {"id": "P003", "nombre": "Accessories Max", "categoria": "Perifericos", "contacto": "info@accmax.com"}
]
with open(os.path.join(data_dir, "proveedores.json"), "w", encoding="utf-8") as f:
    json.dump(proveedores, f, indent=4)

# 2. Crear un Excel: reporte_ventas_Q2.xlsx
ventas_data = {
    "Mes": ["Abril", "Mayo", "Junio"],
    "Ingresos": [45000, 52000, 61000],
    "Unidades_Vendidas": [320, 410, 480],
    "Producto_Estrella": ["Laptop Pro X", "Monitor 4K", "Smartphone Z"]
}
df_ventas = pd.DataFrame(ventas_data)
df_ventas.to_excel(os.path.join(data_dir, "reporte_ventas_Q2.xlsx"), index=False)

# 3. Crear un CSV: registro_asistencia.csv
asistencia_data = {
    "Empleado": ["Ana Garcia", "Carlos Lopez", "Maria Fernanda"],
    "Dias_Trabajados": [22, 20, 22],
    "Horas_Extra": [5, 0, 12],
    "Departamento": ["Ventas", "Soporte", "Administracion"]
}
df_asistencia = pd.DataFrame(asistencia_data)
df_asistencia.to_csv(os.path.join(data_dir, "registro_asistencia.csv"), index=False)

# 4. Crear un TXT: directrices_marketing.txt
marketing_txt = """
Directrices de Marketing - TechStore
1. Tono de comunicacion: Profesional pero accesible, siempre enfocado en los beneficios tecnologicos.
2. Colores de la marca: Azul corporativo (#1E3A8A) y blanco.
3. Promociones: Maximo 2 campanas grandes al ano (Black Friday y Aniversario de la tienda).
4. Redes Sociales: Publicar al menos 3 veces por semana en Instagram y LinkedIn.
"""
with open(os.path.join(data_dir, "directrices_marketing.txt"), "w", encoding="utf-8") as f:
    f.write(marketing_txt)

# 5. Crear un DOCX: manual_empleados.docx
doc = Document()
doc.add_heading('Manual del Empleado - TechStore', 0)
doc.add_heading('1. Cultura de la Empresa', level=1)
doc.add_paragraph('En TechStore valoramos la innovación, el respeto y el compromiso con nuestros clientes.')
doc.add_heading('2. Beneficios', level=1)
doc.add_paragraph('Todos los empleados tienen un bono anual por cumplimiento de metas y seguro medico privado.')
doc.add_heading('3. Codigo de Vestimenta', level=1)
doc.add_paragraph('Smart casual. Se permite ropa cómoda siempre y cuando mantenga un aspecto profesional para la atencion al cliente.')
doc.save(os.path.join(data_dir, "manual_empleados.docx"))

print("¡Archivos de base de conocimiento adicionales creados con éxito!")
